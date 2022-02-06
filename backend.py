import docx
import math

def process_word(filename, save_name, title, prog_code, ver, written_by, edited_by):

    def format_text(given_text, bld = False, itl = False, indent = 12.2):
        """
        For every cell in the row, it will:
        - Set the font to Arial 12 (always)
        - Bold the font if told to (default false)
        - Italic the font if told to (default false)
        - Align in the position given (default left)
        - Indent given, default 12.2 mm
        """
        given_text.runs[0].bold = bld
        given_text.runs[0].itl = itl
        given_text.runs[0].font.name = "Arial"
        given_text.runs[0].font.size = docx.shared.Pt(12)
        given_text.paragraph_format.left_indent = docx.shared.Mm(indent)
        

    def format_table(given_row, bld = False, itl = False, aln = "left"):
        """
        For every cell in the row, it will:
        - Set the font to Arial 12 (always)
        - Bold the font if told to (default false)
        - Italic the font if told to (default false)
        - Align in the position given (default left)
        """

        for cell in given_row:
            # It will always be size 12 Arial
            cell.paragraphs[0].runs[0].font.name = "Arial"
            cell.paragraphs[0].runs[0].font.size = docx.shared.Pt(12)

            if bld:
                cell.paragraphs[0].runs[0].bold = True
            
            if itl:
                cell.paragraphs[0].runs[0].italic = True

            alignments = {
                            "left" : docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT,
                            "centre" : docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER,
                            "right" : docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT,
                        }
            
            cell.paragraphs[0].paragraph_format.alignment = alignments[aln] # docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    document = filename
    output_document = docx.Document()

    script_table = []

    total_time = 0
    counter = 0

    for par in document.paragraphs:
        # Create and format the scene number with leadins 0s
        scene_number = f'{counter:03}'

        # Dealing with word count and text
        text = par.text
        lower_text = text.lower()

        # Deciding if text belongs in visual or audio column
        video = ""
        audio = ""
        time = 2

        if lower_text.split(" ")[0:2] == ["onscreen", "title:"]:
            video = text
            audio = "MUSIC"
        elif lower_text.split(" ")[0:2] == ["chapter", "heading:"]:
            video = text
        else:
            audio = [text]
            word_count = len(text.split(" "))
            time = math.floor(word_count/2)

            if time == 0:
                continue
        
        # Running sum of total time
        total_time += time

        # Add to script table
        script_table.append([scene_number, time, video, audio]) 

        # NB: this is for scene number, not the for loop!
        counter += 1

    # Formatting  the total time
    total_time_formatted = str(math.floor(total_time/60)) + " minutes and " + \
        str(round(((total_time/60) - math.floor(total_time/60))*60 )) + \
            " seconds"
    

    # Setting the size
    section = output_document.sections[0]
    section.page_width = docx.shared.Mm(210)
    section.page_height = docx.shared.Mm(297)
    section.left_margin = docx.shared.Mm(31.7-12.2)
    section.right_margin = docx.shared.Mm(31.7)
    section.top_margin = docx.shared.Mm(22.2)
    section.bottom_margin = docx.shared.Mm(19.0)

    # Creating the header
    header = output_document.sections[0].header
    head = header.paragraphs[0]
    head.text = 'Script Template'
    format_text(head)
    #head.paragraph_format.left_indent = docx.shared.Mm(12.2)

    # Creating the topmatter (uses information provided by the GUI)
    title_of_program = format_text(output_document.add_paragraph("Title of program: " + title), bld=True)
    code_of_program = format_text(output_document.add_paragraph("Program code: " + prog_code), bld=True)
    version_of_program = format_text(output_document.add_paragraph("Version: " + ver), bld=True)
    duration_of_program = format_text(output_document.add_paragraph("Duration: " + total_time_formatted), bld=True)
    writer_of_program = format_text(output_document.add_paragraph("Script Written By: " + written_by), bld=True)
    editor_of_program = format_text(output_document.add_paragraph("Script Edited By: " + edited_by), bld=True)

    
    format_text(output_document.add_paragraph(" "))

    # Creating the table
    document_table = output_document.add_table(rows = 1, cols = 4)
    document_table.style = 'Table Grid'
    document_table.autofit = False 
    document_table.allow_autofit = False

    # Creating the table headers 
    first_row = document_table.rows[0].cells
    first_row[0].text = "SCENE"
    first_row[0].width = docx.shared.Cm(1.91)
    first_row[1].text = "TIME"
    first_row[1].width = docx.shared.Cm(2.01)
    first_row[2].text = "VISION"
    first_row[2].width = docx.shared.Cm(5.03)
    first_row[3].text = "AUDIO"
    first_row[3].width = docx.shared.Cm(8.44)

    # Bolding, italicising, and centring the entire first row
    format_table(first_row, bld=True, itl=True, aln="centre")

    # Populating the document table with entries from "script_table" and applying correct formatting
    row_cntr = 0
    for table_row in script_table:
        row = document_table.add_row().cells
        row[0].text = table_row[0]
        row[0].width = docx.shared.Cm(1.91)
        row[1].text = str(table_row[1])
        row[1].width = docx.shared.Cm(2.01)
        row[2].text = table_row[2]
        row[2].width = docx.shared.Cm(5.03)
        row[3].text = table_row[3]
        row[3].width = docx.shared.Cm(8.44)

        format_table(row[0:2], bld=True, itl=True, aln="centre")
        format_table(row[2:3], bld=True, itl=False, aln="left")

        if row_cntr == 0:
            format_table(row[3::], bld=True, itl=False, aln="left")
        else:
            format_table(row[3::], bld=False, itl=False, aln="left")
        
        row_cntr += 1

    # Creating and formatting the last row
    last_row = document_table.add_row().cells
    last_row[0].text = "Total"
    last_row[0].width = docx.shared.Cm(1.91)
    last_row[1].text = str(total_time)
    last_row[1].width = docx.shared.Cm(2.01)
    last_row[2].text = ""
    last_row[2].width = docx.shared.Cm(5.03)
    last_row[3].text = ""
    last_row[3].width = docx.shared.Cm(8.44)
    format_table(last_row, bld=True, itl=True, aln="centre")

    # Saving the document
    try:
        output_document.save(save_name)
        return True
    except:
        return False